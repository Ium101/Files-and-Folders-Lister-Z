import os
import json
import re
from docx import Document
from docx.shared import Pt

# Language dictionaries
LANGUAGES = {
    "en": {
        "select_language": "Select language / Selecione o idioma:\n1. Português Brasileiro\n2. English\nEnter 1 or 2: ",
        "action_prompt": "What do you want to do?\n1. Create a list of files and folders\n2. Create a folder structure from a JSON file\nEnter your choice (1/2): ",
        "invalid_action": "Invalid choice. Please enter 1 or 2.",
        "enter_directory": "Enter the directory to list: ",
        "error_directory": "Error: The directory '{dir}' does not exist. Please enter a valid folder.",
        "hide_hidden": "Do you want to hide the files desktop.ini, thumbs.db, ._.ds_store, .ds_store, .gitignore, and .gitkeep? (y/yes/n/no): ",
        "filter_input": "Enter sub-folder names or keywords to filter (comma-separated, case-insensitive, substring match), or leave blank to include all: ",
        "output_mode": "Do you want to generate the output as DOCX (A), TXT (B), or JSON (C)? ",
        "invalid_mode": "Invalid input. Please enter DOCX/A or TXT/B, or JSON/C.",
        "choose_option": "Choose listing option:",
        "option_1": "  1. List folders and files",
        "option_2": "  2. List only folders",
        "option_3": "  3. List only files",
        "enter_choice": "Enter your choice (1/2/3): ",
        "invalid_choice": "Invalid choice. Please enter 1, 2, or 3.",
        "invalid_number": "Invalid input. Please enter a number (1, 2, or 3).",
        "list_success": "List generated successfully: {path}",
        "list_generated": "The List has been generated",
        "docx_failed": "DOCX generation failed: {err}. Falling back to TXT mode.",
        "json_exported": "JSON database exported: {path}",
        "enter_json_path": "Enter the path to the JSON file: ",
        "error_json_path": "Error: The file '{path}' does not exist.",
        "error_json_format": "Error: The file '{path}' is not a valid JSON file. Please check the format.",
        "enter_base_dir": "Enter the base directory where the folders will be created: ",
        "folders_created_successfully": "Folder structure created successfully!",
        "error_creating_folders": "An error occurred while creating folders: {e}",
        "credits": "Made by User Ium101 from GitHub",
        "press_any_button": "Press any button to exit"
    },
    "pt": {
        "select_language": "Selecione o idioma / Select language:\n1. Português Brasileiro\n2. English\nDigite 1 ou 2: ",
        "action_prompt": "O que você deseja fazer?\n1. Criar uma lista de arquivos e pastas\n2. Criar uma estrutura de pastas a partir de um arquivo JSON\nDigite sua escolha (1/2): ",
        "invalid_action": "Escolha inválida. Por favor, insira 1 ou 2.",
        "enter_directory": "Digite o diretório para listar: ",
        "error_directory": "Erro: O diretório '{dir}' não existe. Por favor, insira uma pasta válida.",
        "hide_hidden": "Deseja ocultar os arquivos desktop.ini, thumbs.db, ._.ds_store, .ds_store, .gitignore, e .gitkeep? (s/sim/n/não): ",
        "filter_input": "Digite nomes de subpastas ou palavras-chave para filtrar (separados por vírgula, sem diferenciar maiúsculas/minúsculas), ou deixe em branco para incluir todas: ",
        "output_mode": "Deseja gerar a saída como DOCX (A), TXT (B) ou JSON (C)? ",
        "invalid_mode": "Entrada inválida. Por favor, insira DOCX/A, TXT/B ou JSON/C.",
        "choose_option": "Escolha a opção de listagem:",
        "option_1": "  1. Listar pastas e arquivos",
        "option_2": "  2. Listar apenas pastas",
        "option_3": "  3. Listar apenas arquivos",
        "enter_choice": "Digite sua escolha (1/2/3): ",
        "invalid_choice": "Escolha inválida. Por favor, insira 1, 2 ou 3.",
        "invalid_number": "Entrada inválida. Por favor, insira um número (1, 2 ou 3).",
        "list_success": "Lista gerada com sucesso: {path}",
        "list_generated": "A lista foi gerada",
        "docx_failed": "Falha ao gerar DOCX: {err}. Gerando TXT em vez disso.",
        "json_exported": "Banco de dados JSON exportado: {path}",
        "enter_json_path": "Digite o caminho para o arquivo JSON: ",
        "error_json_path": "Erro: O arquivo '{path}' não existe.",
        "error_json_format": "Erro: O arquivo '{path}' não é um arquivo JSON válido. Por favor, verifique o formato.",
        "enter_base_dir": "Digite o diretório base onde as pastas serão criadas: ",
        "folders_created_successfully": "Estrutura de pastas criada com sucesso!",
        "error_creating_folders": "Ocorreu um erro ao criar as pastas: {e}",
        "credits": "Feito pelo Usuário Ium101 do GitHub",
        "press_any_button": "Pressione qualquer botão para sair"
    }
}

def get_lang():
    while True:
        lang_choice = input(LANGUAGES["en"]["select_language"])
        if lang_choice.strip() == "1":
            return "pt"
        elif lang_choice.strip() == "2":
            return "en"
        else:
            print("Invalid input. Please enter 1 or 2. / Entrada inválida. Por favor, insira 1 ou 2.")

def list_files_and_folders(directory, mode="B", list_option=1, recursive=False, specific_subfolders=None, ignore_hidden=False, L=None, lang='en'):
    folder_name = os.path.basename(os.path.normpath(directory))
    disk_letter = os.path.splitdrive(os.path.abspath(directory))[0].replace(":", "")
    output_file_ext = 'docx' if mode.upper() == 'A' else ('json' if mode.upper() == 'C' else 'txt')
    output_file_path = os.path.join(directory, f"{folder_name} ({disk_letter}).{output_file_ext}")
    folders = []
    files = []
    for entry in os.scandir(directory):
        if ignore_hidden and is_hidden_file(entry):
            continue
        if entry.is_dir():
            if not specific_subfolders:
                folders.append(entry.path)
            else:
                entry_name_clean = re.sub(r'\W+', '', entry.name).lower()
                for sub in specific_subfolders:
                    sub_clean = re.sub(r'\W+', '', sub).lower()
                    if sub_clean == entry_name_clean:
                        folders.append(entry.path)
                        break
        elif entry.is_file():
            files.append(entry.path)

    credits_str = L["credits"]

    if mode.upper() == "A":  # DOCX
        try:
            doc = Document()
            doc.add_heading(f"{folder_name}", level=1)
            if list_option in [1, 2]:
                for folder in folders:
                    write_folder_structure_docx(doc, folder, list_option=list_option)
            if list_option == 1:
                for file in files:
                    p = doc.add_paragraph()
                    p.add_run("• ")
                    base, ext = os.path.splitext(os.path.basename(file))
                    p.add_run(base)
                    p.add_run(ext)

            # Add credits at the end (small size)
            p_credits = doc.add_paragraph()
            run = p_credits.add_run(credits_str())
            try:
                run.font.size = Pt(8)
            except Exception:
                # If font sizing fails for any reason, ignore and continue
                pass

            doc.save(output_file_path)
            print(L["list_success"].format(path=output_file_path))
            print(L["list_generated"])
        except Exception as e:
            # Fallback to TXT if DOCX fails; include credits in TXT
            print(L["docx_failed"].format(err=e))
            output_file_path = os.path.join(directory, f"{folder_name} ({disk_letter}).txt")
            with open(output_file_path, "w", encoding="utf-8") as txt_file:
                txt_file.write(f"{folder_name}\n\n")
                if list_option in [1, 2]:
                    for folder in folders:
                        write_folder_structure_txt(txt_file, folder)
                if list_option in [1, 3]:
                    for file in files:
                        base, ext = os.path.splitext(os.path.basename(file))
                        txt_file.write(f"• {base}{ext}\n")
                txt_file.write(f"\n{credits_str}\n")
            print(L["list_success"].format(path=output_file_path))
            print(L["list_generated"])

    elif mode.upper() == "C":  # JSON
        def folder_to_dict(folder):
            d = {"folder": os.path.basename(folder), "files": [], "subfolders": []}
            for entry in os.scandir(folder):
                if entry.is_file():
                    d["files"].append(entry.name)
                elif entry.is_dir():
                    d["subfolders"].append(folder_to_dict(entry.path))
            return d

        db = {
            "root": folder_name,
            "files": [os.path.basename(f) for f in files] if list_option in [1, 3] else [],
            "folders": [],
            "credits": credits_str
        }
        for folder in folders:
            db["folders"].append(folder_to_dict(folder))

        with open(output_file_path, "w", encoding="utf-8") as json_file:
            json.dump(db, json_file, indent=2, ensure_ascii=False)
        print(L["json_exported"].format(path=output_file_path))

    else:  # TXT
        with open(output_file_path, "w", encoding="utf-8") as txt_file:
            txt_file.write(f"{folder_name}\n\n")
            if list_option in [1, 2]:
                for folder in folders:
                    write_folder_structure_txt(txt_file, folder)
            if list_option in [1, 3]:
                for file in files:
                    base, ext = os.path.splitext(os.path.basename(file))
                    txt_file.write(f"• {base}{ext}\n")
            txt_file.write(f"\n{credits_str}\n")
        print(L["list_success"].format(path=output_file_path))
        print(L["list_generated"])

def write_folder_structure_docx(doc, folder, indent=0, list_option=1):
    p = doc.add_paragraph("    " * indent)
    p.add_run("• ")
    run = p.add_run(os.path.basename(folder))
    run.bold = True
    entries = sorted(os.scandir(folder), key=lambda e: (not e.is_dir(), e.name.lower()))
    for entry in entries:
        if entry.is_dir():
            write_folder_structure_docx(doc, entry.path, indent + 1, list_option)
        elif entry.is_file() and list_option == 1:
            sub_p = doc.add_paragraph("    " * (indent + 1))
            base, ext = os.path.splitext(entry.name)
            sub_run = sub_p.add_run(base)
            sub_run.italic = True
            sub_p.add_run(ext)

def write_folder_structure_txt(txt_file, folder, indent=0, list_option=1):
    txt_file.write(f"{'    ' * indent}• {os.path.basename(folder)}\n")
    entries = sorted(os.scandir(folder), key=lambda e: (not e.is_dir(), e.name.lower()))
    for entry in entries:
        if entry.is_dir():
            write_folder_structure_txt(txt_file, entry.path, indent + 1, list_option)
        elif entry.is_file() and list_option == 1:
            base, ext = os.path.splitext(entry.name)
            txt_file.write(f"{'    ' * (indent + 1)}{base}{ext}\n")

def is_hidden_file(entry):
    hidden_names = {"desktop.ini", "thumbs.db", "._.ds_store", ".ds_store", ".gitignore", ".gitkeep"}
    if entry.name.startswith('.') or entry.name.lower() in hidden_names:
        return True
    try:
        import ctypes
        attrs = ctypes.windll.kernel32.GetFileAttributesW(str(entry.path))
        if attrs != -1 and attrs & 2:
            return True
    except Exception:
        pass
    return False

def create_folders_from_json_cli(json_path, base_dir, L):
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except json.JSONDecodeError:
        print(L["error_json_format"].format(path=json_path))
        return False
    except Exception as e:
        print(L["error_creating_folders"].format(e=e))
        return False

    try:
        def create_structure(d, parent):
            folder_name = d.get("folder", d.get("root", ""))
            if not folder_name:
                return
            folder_path = os.path.join(parent, folder_name)
            os.makedirs(folder_path, exist_ok=True)
            for subfolder in d.get("subfolders", []):
                create_structure(subfolder, folder_path)

        if "root" in data:
            root_dict = {"folder": data["root"], "subfolders": data.get("folders", [])}
            create_structure(root_dict, base_dir)
        else:
            create_structure(data, base_dir)

        print(L["folders_created_successfully"])
        return True
    except Exception as e:
        print(L["error_creating_folders"].format(e=e))
        return False

def run_create_list(L, lang):
    while True:
        directory_to_list = input(L["enter_directory"])
        if not os.path.isdir(directory_to_list):
            print(L["error_directory"].format(dir=directory_to_list))
        else:
            break
    hide_hidden_input = input(L["hide_hidden"]).strip().lower()
    ignore_hidden = hide_hidden_input in (["yes", "y"] if lang == "en" else ["sim", "s"])
    filter_input = input(L["filter_input"]).strip()
    specific_subfolders = [folder.strip() for folder in filter_input.split(",")] if filter_input else None

    while True:
        mode = input(L["output_mode"]).strip().lower()
        if mode in ("a", "docx"): mode = "A"; break
        if mode in ("b", "txt"): mode = "B"; break
        if mode in ("c", "json"): mode = "C"; break
        print(L["invalid_mode"])

    while True:
        try:
            print(L["choose_option"])
            print(L["option_1"])
            print(L["option_2"])
            print(L["option_3"])
            choice = input(L["enter_choice"])
            list_option = int(choice)
            if list_option in [1, 2, 3]:
                break
            else:
                print(L["invalid_choice"])
        except ValueError:
            print(L["invalid_number"])

    list_files_and_folders(directory_to_list, mode=mode, list_option=list_option, recursive=True, specific_subfolders=specific_subfolders, ignore_hidden=ignore_hidden, L=L, lang=lang)

def run_create_folders(L):
    while True:
        json_path = input(L["enter_json_path"]).strip()
        if os.path.isfile(json_path):
            break
        else:
            print(L["error_json_path"].format(path=json_path))

    while True:
        base_dir = input(L["enter_base_dir"]).strip()
        if os.path.isdir(base_dir):
            break
        else:
            print(L["error_directory"].format(dir=base_dir))

    create_folders_from_json_cli(json_path, base_dir, L)

if __name__ == "__main__":
    lang = get_lang()
    L = LANGUAGES[lang]

    while True:
        action = input(L["action_prompt"])
        if action.strip() == "1":
            run_create_list(L, lang)
            break
        elif action.strip() == "2":
            run_create_folders(L)
            break
        else:
            print(L["invalid_action"])

    input(L["press_any_button"])