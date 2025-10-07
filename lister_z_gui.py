import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from docx import Document
from docx.shared import Pt
import json

LANGUAGES = {
    "en": {
        "title": "Lister Z",
        "create_list": "Create a list of files and folders (English)",
        "create_folders": "Create folder structure from JSON (English)",
        "select_dir": "Select a folder to list.",
        "error_dir": "The directory '{directory}' does not exist. Please select a valid folder.",
        "mode": "Do you want to generate the output as DOCX (A), TXT (B), or JSON (C)?",
        "invalid_mode": "Invalid input. Please enter DOCX/A or TXT/B, or JSON/C.",
        "list_option": "Choose listing option:\n1. Folders and files\n2. Only folders\n3. Only files",
        "filter": "Enter sub-folder names or keywords to filter (comma-separated, case-insensitive), or click OK to include all:",
        "hide_hidden": "Do you want to hide system and hidden files (like desktop.ini, .ds_store)?",
        "success": "List generated successfully: {path}",
        "json_success": "JSON database exported: {path}",
        "docx_error": "DOCX generation failed: {err}",
        "credits": "Made by User Ium101 from GitHub",
        "select_json": "Select the JSON file for folder creation",
        "no_json_selected": "Operation cancelled: No JSON file was selected.",
        "invalid_json_format": "The selepyinstaller --onefile --windowed lister_z_gui.py ted file is not a valid JSON file. Please choose a file with the correct format.",
        "select_base_dir": "Select the base folder to create the structure in",
        "no_base_dir_selected": "Operation cancelled: No base folder was selected.",
        "folders_created": "Folder structure created successfully!",
        "error_creating": "Error creating folders",
        "error": "Error"
    },
    "pt": {
        "title": "Lister Z",
        "create_list": "Criar uma lista de arquivos e pastas (Português)",
        "create_folders": "Criar estrutura de pastas a partir de JSON (Português)",
        "select_dir": "Selecione uma pasta para listar.",
        "error_dir": "O diretório '{directory}' não existe. Por favor, selecione uma pasta válida.",
        "mode": "Deseja gerar a saída como DOCX (A), TXT (B) ou JSON (C)?",
        "invalid_mode": "Entrada inválida. Por favor, insira DOCX/A, TXT/B ou JSON/C.",
        "list_option": "Escolha a opção de listagem:\n1. Pastas e arquivos\n2. Apenas pastas\n3. Apenas arquivos",
        "filter": "Digite nomes de subpastas ou palavras-chave para filtrar (separados por vírgula, sem diferenciar maiúsculas/minúsculas), ou clique em OK para incluir todas:",
        "hide_hidden": "Deseja ocultar arquivos de sistema e ocultos (como desktop.ini, .ds_store)?",
        "success": "Lista gerada com sucesso: {path}",
        "json_success": "Banco de dados JSON exportado: {path}",
        "docx_error": "Falha ao gerar DOCX: {err}",
        "credits": "Feito pelo Usuário Ium101 do GitHub",
        "select_json": "Selecione o arquivo JSON para a criação de pastas",
        "no_json_selected": "Operação cancelada: Nenhum arquivo JSON foi selecionado.",
        "invalid_json_format": "O arquivo selecionado não é um JSON válido. Por favor, escolha um arquivo com o formato correto.",
        "select_base_dir": "Selecione a pasta base para criar a estrutura",
        "no_base_dir_selected": "Operação cancelada: Nenhuma pasta base foi selecionada.",
        "folders_created": "Estrutura de pastas criada com sucesso!",
        "error_creating": "Erro ao criar pastas",
        "error": "Erro"
    }
}

def is_hidden_file(entry):
    # Cross-platform: check dotfiles and a set of known hidden names, try Windows attributes when available.
    hidden_names = {"desktop.ini", "thumbs.db", "._.ds_store", ".ds_store", ".gitignore", ".gitkeep"}
    if entry.name.startswith('.') or entry.name.lower() in hidden_names:
        return True
    try:
        import ctypes
        attrs = ctypes.windll.kernel32.GetFileAttributesW(str(entry.path))
        if attrs != -1 and attrs & 2:  # FILE_ATTRIBUTE_HIDDEN
            return True
    except Exception:
        pass
    return False

def list_files_and_folders(directory, mode="B", list_option=1, recursive=True, specific_subfolders=None, ignore_hidden=False, L=None, parent=None):
    folder_name = os.path.basename(os.path.normpath(directory))
    disk_letter = os.path.splitdrive(os.path.abspath(directory))[0].replace(":", "")
    output_filename_base = f"{folder_name} ({disk_letter})"
    folders, files = [], []

    for entry in os.scandir(directory):
        if ignore_hidden and is_hidden_file(entry):
            continue
        if entry.is_dir():
            if not specific_subfolders or any(re.sub(r'\W+', '', sub).lower() in re.sub(r'\W+', '', entry.name).lower() for sub in specific_subfolders):
                folders.append(entry.path)
        elif entry.is_file():
            files.append(entry.path)

    output_file_ext = 'docx' if mode.upper() == 'A' else ('json' if mode.upper() == 'C' else 'txt')
    output_file_path = os.path.join(directory, f"{output_filename_base}.{output_file_ext}")
    credits_str = L["credits"]

    if mode.upper() == "A":
        try:
            doc = Document()
            doc.add_heading(f"{folder_name}", level=1)

            if list_option in [1, 2]:
                for folder in folders:
                    write_folder_structure_docx(doc, folder, list_option=list_option)

            if list_option == 1:
                for file in files:
                    p = doc.add_paragraph("• ")
                    base, ext = os.path.splitext(os.path.basename(file))
                    p.add_run(base).italic = True
                    p.add_run(ext)

            # Add credits in small font size (8pt)
            p_credits = doc.add_paragraph()
            run = p_credits.add_run(credits_str)
            try:
                run.font.size = Pt(8)
            except Exception:
                # If setting font size fails for any reason, ignore it but still save the credits text
                pass

            doc.save(output_file_path)
            messagebox.showinfo(L["title"], L["success"].format(path=output_file_path), parent=parent)
        except Exception as e:
            messagebox.showerror(L["error"], L["docx_error"].format(err=e), parent=parent)

    elif mode.upper() == "C":
        def folder_to_dict(path):
            d = {"folder": os.path.basename(path), "files": [], "subfolders": []}
            for entry in os.scandir(path):
                if entry.is_file():
                    d["files"].append(entry.name)
                elif entry.is_dir():
                    d["subfolders"].append(folder_to_dict(entry.path))
            return d

        db = {
            "root": folder_name,
            "files": [os.path.basename(f) for f in files if list_option in [1, 3]],
            "folders": [folder_to_dict(f) for f in folders],
            "credits": credits_str
        }
        with open(output_file_path, "w", encoding="utf-8") as json_file:
            json.dump(db, json_file, indent=2, ensure_ascii=False)
        messagebox.showinfo(L["title"], L["json_success"].format(path=output_file_path), parent=parent)

    else:  # TXT
        with open(output_file_path, "w", encoding="utf-8") as txt_file:
            txt_file.write(f"{folder_name}\n\n")
            if list_option in [1, 2]:
                for folder in folders:
                    write_folder_structure_txt(txt_file, folder, list_option=list_option)
            if list_option in [1, 3]:
                for file in files:
                    txt_file.write(f"• {os.path.basename(file)}\n")
            # Append credits at the end (normal case)
            txt_file.write(f"\n{credits_str}\n")
        messagebox.showinfo(L["title"], L["success"].format(path=output_file_path), parent=parent)

def write_folder_structure_docx(doc, folder, indent=0, list_option=1):
    p = doc.add_paragraph("    " * indent + "• ")
    p.add_run(os.path.basename(folder)).bold = True
    entries = sorted(os.scandir(folder), key=lambda e: (not e.is_dir(), e.name.lower()))
    for entry in entries:
        if entry.is_dir():
            write_folder_structure_docx(doc, entry.path, indent + 1, list_option)
        elif entry.is_file() and list_option == 1:
            sub_p = doc.add_paragraph("    " * (indent + 1))
            base, ext = os.path.splitext(entry.name)
            sub_p.add_run(base).italic = True
            sub_p.add_run(ext)

def write_folder_structure_txt(txt_file, folder, indent=0, list_option=1):
    txt_file.write(f"{'    ' * indent}• {os.path.basename(folder)}\n")
    entries = sorted(os.scandir(folder), key=lambda e: (not e.is_dir(), e.name.lower()))
    for entry in entries:
        if entry.is_dir():
            write_folder_structure_txt(txt_file, entry.path, indent + 1, list_option)
        elif entry.is_file() and list_option == 1:
            txt_file.write(f"{'    ' * (indent + 1)}{os.path.basename(entry.name)}\n")

def create_folders_from_json_gui(root, lang_code):
    L = LANGUAGES[lang_code]
    root.withdraw()
    while True:
        json_path = filedialog.askopenfilename(title=L["select_json"], filetypes=[("JSON files", "*.json")], parent=root)
        if not json_path:
            root.deiconify()
            return
        try:
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            break
        except json.JSONDecodeError:
            messagebox.showerror(L["error"], L["invalid_json_format"], parent=root)
        except Exception as e:
            messagebox.showerror(L["error"], f"{L['error_creating']}: {e}", parent=root)
            root.deiconify()
            return

    base_dir = filedialog.askdirectory(title=L["select_base_dir"], parent=root)
    if not base_dir:
        root.deiconify()
        return

    try:
        def create_structure(d, parent):
            folder_name = d.get("folder", d.get("root", ""))
            if folder_name:
                folder_path = os.path.join(parent, folder_name)
                os.makedirs(folder_path, exist_ok=True)
                for subfolder in d.get("subfolders", []):
                    create_structure(subfolder, folder_path)

        if "root" in data:
            create_structure({"folder": data["root"], "subfolders": data.get("folders", [])}, base_dir)
        else:
            create_structure(data, base_dir)

        messagebox.showinfo(L["title"], L["folders_created"], parent=root)
    except Exception as e:
        messagebox.showerror(L["error"], f"{L['error_creating']}: {e}", parent=root)

    root.deiconify()

def run_lister(root, lang_code):
    L = LANGUAGES[lang_code]
    root.withdraw()
    directory = filedialog.askdirectory(title=L["select_dir"], parent=root)
    if not directory:
        root.deiconify()
        return

    mode = simpledialog.askstring(L["title"], L["mode"], parent=root)
    if mode is None:
        root.deiconify()
        return
    mode = mode.strip().lower()
    if mode in ["a", "docx"]:
        mode = "A"
    elif mode in ["b", "txt"]:
        mode = "B"
    elif mode in ["c", "json"]:
        mode = "C"
    else:
        messagebox.showerror(L["title"], L["invalid_mode"], parent=root)
        root.deiconify()
        return

    list_option = simpledialog.askinteger(L["title"], L["list_option"], minvalue=1, maxvalue=3, parent=root)
    if list_option is None:
        root.deiconify()
        return

    filter_input = simpledialog.askstring(L["title"], L["filter"], parent=root)
    if filter_input is None:
        root.deiconify()
        return
    specific_subfolders = [f.strip() for f in filter_input.split(",")] if filter_input else None

    ignore_hidden = messagebox.askyesno(L["title"], L["hide_hidden"], parent=root)
    list_files_and_folders(directory, mode=mode, list_option=list_option, recursive=True, specific_subfolders=specific_subfolders, ignore_hidden=ignore_hidden, L=L, parent=root)
    root.deiconify()

def run_gui():
    root = tk.Tk()
    root.title("Lister Z")
    root.geometry("500x250")

    L_en = LANGUAGES["en"]
    L_pt = LANGUAGES["pt"]

    list_en_btn = tk.Button(root, text=L_en["create_list"], command=lambda: run_lister(root, "en"), height=2, width=50)
    list_en_btn.pack(pady=5)

    list_pt_btn = tk.Button(root, text=L_pt["create_list"], command=lambda: run_lister(root, "pt"), height=2, width=50)
    list_pt_btn.pack(pady=5)

    create_en_btn = tk.Button(root, text=L_en["create_folders"], command=lambda: create_folders_from_json_gui(root, "en"), height=2, width=50)
    create_en_btn.pack(pady=5)

    create_pt_btn = tk.Button(root, text=L_pt["create_folders"], command=lambda: create_folders_from_json_gui(root, "pt"), height=2, width=50)
    create_pt_btn.pack(pady=5)

    credits_frame = tk.Frame(root)
    credits_frame.pack(side="bottom", pady=10)
    credits_lbl = tk.Label(credits_frame, text=f'{L_pt["credits"]} / {L_en["credits"]}', font=("Arial", 8))
    credits_lbl.pack()

    root.mainloop()

if __name__ == "__main__":
    run_gui()